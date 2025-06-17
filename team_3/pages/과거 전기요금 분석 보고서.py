import streamlit as st
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.express as px
from datetime import datetime, timedelta
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from matplotlib import rcParams

# 페이지 설정
st.set_page_config(page_title="통합 전력 분석", layout="wide")

# matplotlib 한글 폰트 설정
plt.rcParams['font.family'] = ['Malgun Gothic', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# CSS 스타일링
st.markdown("""
<style>
    .stApp { background-color: #fafafa; }
    .metric-card { background-color: white; padding: 1rem; border-radius: 10px; border: 1px solid #e0e0e0; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
    .comparison-table { background-color: white; border-radius: 10px; padding: 1rem; box-shadow: 0 2px 8px rgba(0,0,0,0.08); border: 1px solid #e0e0e0; }
    .header-style { background: white; color: #2c3e50; padding: 2rem; border-radius: 12px; text-align: center; margin-bottom: 2rem; box-shadow: 0 4px 12px rgba(0,0,0,0.1); border-left: 4px solid #3498db; }
    .section-header { background: white; color: #2c3e50; padding: 1.2rem; border-radius: 12px; text-align: center; margin: 1rem 0; box-shadow: 0 2px 8px rgba(0,0,0,0.08); border-left: 4px solid #2980b9; }
    .time-period-card { background: white; border-radius: 12px; padding: 20px; color: #2c3e50; text-align: center; box-shadow: 0 3px 10px rgba(0,0,0,0.1); margin-bottom: 15px; transition: all 0.3s ease; border: 1px solid #e0e0e0; }
    .time-period-card:hover { transform: translateY(-2px); box-shadow: 0 6px 20px rgba(0,0,0,0.15); }
    .daytime-card { border-left: 4px solid #f39c12; }
    .nighttime-card { border-left: 4px solid #34495e; }
    .card-title { font-size: 1.0em; font-weight: 600; margin-bottom: 8px; color: #34495e; }
    .card-value { font-size: 1.8em; font-weight: 700; margin: 8px 0; color: #2c3e50; }
    .card-change { font-size: 0.9em; margin: 4px 0; font-weight: 500; color: #7f8c8d; }
    .card-period { font-size: 0.75em; color: #95a5a6; margin-top: 8px; }
    .traffic-light { font-size: 1.2em; margin-right: 8px; }
    .main-metrics-card { background: white; border-radius: 16px; padding: 24px; color: #2c3e50; margin: 20px 0; box-shadow: 0 4px 16px rgba(0,0,0,0.08); border: 1px solid #e8ecef; }
    .metrics-grid { display: grid; grid-template-columns: 1fr 1fr 1fr 1fr; gap: 20px; margin-top: 20px; }
    .metric-item { background: #f8f9fa; border-radius: 12px; padding: 16px; text-align: center; border: 1px solid #e9ecef; transition: all 0.2s ease; }
    .metric-item:hover { background: #e9ecef; }
    .metric-label { font-size: 0.9em; color: #6c757d; margin-bottom: 8px; font-weight: 500; }
    .metric-value { font-size: 1.5em; font-weight: 700; color: #2c3e50; }
</style>
""", unsafe_allow_html=True)

# ========== 1. 데이터 로드 함수 ==========
@st.cache_data
def load_data():
    """데이터 로드 및 전처리"""
    try:
        df = pd.read_csv("./data/train.csv")
        df["측정일시"] = pd.to_datetime(df["측정일시"])
        df["날짜"] = df["측정일시"].dt.date
        df["시간"] = df["측정일시"].dt.hour
        df["월"] = df["측정일시"].dt.month
        df["일"] = df["측정일시"].dt.day
        df["년월"] = df["측정일시"].dt.to_period("M")
        return df
    except Exception as e:
        st.error(f"데이터 로드 중 오류 발생: {e}")
        return None

# ========== 2. 차트 생성 함수들 ==========
def create_matplotlib_chart(data, chart_type="line", title="Chart", xlabel="X", ylabel="Y", figsize=(10, 6)):
    """matplotlib로 간단한 차트 생성"""
    fig, ax = plt.subplots(figsize=figsize)
    fig.patch.set_facecolor('white')
    
    if chart_type == "line" and len(data.columns) >= 2:
        x_data = data.iloc[:, 0]
        y_data = data.iloc[:, 1]
        ax.plot(x_data, y_data, marker='o', linewidth=2, markersize=6, color='#1f77b4')
    
    elif chart_type == "bar" and len(data.columns) >= 2:
        x_data = data.iloc[:, 0]
        y_data = data.iloc[:, 1]
        bars = ax.bar(x_data, y_data, color=['#1f77b4', '#ff7f0e', '#2ca02c'])
        
        # 막대 위에 값 표시
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{height:,.0f}', ha='center', va='bottom')
    
    elif chart_type == "pie" and len(data.columns) >= 2:
        labels = data.iloc[:, 0]
        sizes = data.iloc[:, 1]
        colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99']
        wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%', 
                                         colors=colors, startangle=90)
        ax.axis('equal')
    
    ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
    ax.set_xlabel(xlabel, fontsize=12)
    ax.set_ylabel(ylabel, fontsize=12)
    ax.grid(True, alpha=0.3)
    
    plt.tight_layout()
    
    # 이미지로 변환
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer

def create_dual_axis_chart(df, x_col, y1_col, y2_col, title, x_title, y1_title, y2_title, add_time_zones=False):
    """듀얼 축 차트 생성"""
    fig = make_subplots(specs=[[{"secondary_y": True}]])

    if add_time_zones and x_title == "시간":
        fig.add_vrect(x0=-0.5, x1=5.5, fillcolor="rgba(25, 25, 112, 0.1)", layer="below", line_width=0,
                     annotation_text="야간 (00-05시)", annotation_position="top left")
        fig.add_vrect(x0=5.5, x1=17.5, fillcolor="rgba(255, 255, 0, 0.05)", layer="below", line_width=0,
                     annotation_text="주간 (06-17시)", annotation_position="top")
        fig.add_vrect(x0=17.5, x1=21.5, fillcolor="rgba(128, 0, 128, 0.08)", layer="below", line_width=0,
                     annotation_text="저녁 (18-21시)", annotation_position="top right")
        fig.add_vrect(x0=21.5, x1=23.5, fillcolor="rgba(25, 25, 112, 0.1)", layer="below", line_width=0,
                     annotation_text="야간 (22-23시)", annotation_position="top right")

    fig.add_trace(go.Scatter(x=df[x_col], y=df[y1_col], name=y1_title, line=dict(color="#1f77b4", width=3),
                           mode="lines+markers", marker=dict(size=6)), secondary_y=False)
    fig.add_trace(go.Scatter(x=df[x_col], y=df[y2_col], name=y2_title, line=dict(color="#ff7f0e", width=3),
                           mode="lines+markers", marker=dict(size=6)), secondary_y=True)

    fig.update_xaxes(title_text=x_title)
    fig.update_yaxes(title_text=y1_title, secondary_y=False, title_font_color="#1f77b4")
    fig.update_yaxes(title_text=y2_title, secondary_y=True, title_font_color="#ff7f0e")
    fig.update_layout(title=title, hovermode="x unified", template="plotly_white", height=500,
                     font=dict(family="맑은 고딕"))
    return fig

def create_hourly_stack_chart(df):
    """시간별 스택 차트 생성"""
    hourly_worktype = df.groupby(["시간", "작업유형"])["전기요금(원)"].sum().unstack(fill_value=0)
    colors = {"Light_Load": "rgba(76, 175, 80, 0.7)", "Medium_Load": "rgba(255, 152, 0, 0.7)", "Maximum_Load": "rgba(244, 67, 54, 0.7)"}

    fig = go.Figure()
    for work_type in ["Light_Load", "Medium_Load", "Maximum_Load"]:
        if work_type in hourly_worktype.columns:
            fig.add_trace(go.Bar(name=work_type, x=hourly_worktype.index, y=hourly_worktype[work_type],
                               marker_color=colors.get(work_type)))

    fig.update_layout(barmode="stack", title="시간대별 작업유형별 전기요금 현황", xaxis_title="시간 (Hour)",
                     yaxis_title="전기요금 (원)", height=500, plot_bgcolor="white", paper_bgcolor="white",
                     font=dict(family="맑은 고딕"))
    return fig

def create_concentric_donut_chart(df):
    """도넛 차트 생성"""
    worktype_mwh = df.groupby("작업유형")["전력사용량(kWh)"].sum() / 1000
    total_mwh = worktype_mwh.sum()

    chart_data_map = {"Light_Load": {"name": "경부하", "color": "#4CAF50"},
                     "Medium_Load": {"name": "중간부하", "color": "#FF9800"},
                     "Maximum_Load": {"name": "최대부하", "color": "#F44336"}}

    labels, values, colors = [], [], []
    for work_type, data in chart_data_map.items():
        if work_type in worktype_mwh.index:
            labels.append(data["name"])
            values.append(worktype_mwh[work_type])
            colors.append(data["color"])

    fig = go.Figure(data=[go.Pie(labels=labels, values=values, hole=0.55, marker=dict(colors=colors))])
    fig.update_layout(title="부하대별 에너지 사용량 분포", height=500,
                     annotations=[dict(text=f"총 {total_mwh:,.1f} MWh", x=0.5, y=0.5, showarrow=False)],
                     font=dict(family="맑은 고딕"))
    return fig

# ========== 3. 카드 및 테이블 생성 함수들 ==========
def create_main_metrics_card(summary_data, period_label):
    """주요 지표 카드 생성"""
    if summary_data.empty:
        return ""
    
    total_kwh = summary_data["전력사용량(kWh)"].sum()
    total_cost = summary_data["전기요금(원)"].sum()
    total_carbon = summary_data["탄소배출량(tCO2)"].sum()
    avg_price = total_cost / total_kwh if total_kwh > 0 else 0
    
    card_html = f"""
    <div class="main-metrics-card">
        <h3 style="text-align: center; margin-bottom: 15px; color: #333;"> {period_label} 주요 지표</h3>
        <div class="metrics-grid">
            <div class="metric-item">
                <div class="metric-label">전력사용량</div>
                <div class="metric-value">{total_kwh:,.1f} kWh</div>
            </div>
            <div class="metric-item">
                <div class="metric-label">전기요금</div>
                <div class="metric-value">{total_cost:,.0f} 원</div>
            </div>
            <div class="metric-item">
                <div class="metric-label">평균 단가</div>
                <div class="metric-value">{avg_price:.1f} 원/kWh</div>
            </div>
            <div class="metric-item">
                <div class="metric-label">탄소배출량</div>
                <div class="metric-value">{total_carbon:.2f} tCO2</div>
            </div>
        </div>
    </div>
    """
    return card_html

def calculate_kepco_rate_impact(pf_value, time_period):
    """한전 요금 영향 계산"""
    if time_period == "daytime":
        adjusted_pf = max(60, min(95, pf_value))
        if adjusted_pf >= 90:
            rate_impact = -(adjusted_pf - 90) * 0.5
        else:
            rate_impact = (90 - adjusted_pf) * 0.5
    else:
        if pf_value <= 0:
            adjusted_pf = 100
        else:
            adjusted_pf = max(60, pf_value)
        if adjusted_pf >= 95:
            rate_impact = 0
        else:
            rate_impact = (95 - adjusted_pf) * 0.5
    return rate_impact

def get_traffic_light_and_message(current_pf, previous_pf, time_period):
    """신호등 및 메시지 생성"""
    current_impact = calculate_kepco_rate_impact(current_pf, time_period)
    previous_impact = calculate_kepco_rate_impact(previous_pf, time_period)
    rate_difference = current_impact - previous_impact
    
    if abs(rate_difference) < 0.1:
        traffic_light = "🟡"
        message = "전일과 동일"
    elif rate_difference > 0:
        traffic_light = "🔴"
        message = f"전일대비 +{rate_difference:.1f}% 더 냄"
    else:
        traffic_light = "🟢"
        message = f"전일대비 {rate_difference:.1f}% 덜 냄"
    return traffic_light, message

def create_simple_power_factor_card(period_name, icon, current_pf, previous_pf, time_period, card_class):
    """역률 카드 생성"""
    traffic_light, message = get_traffic_light_and_message(current_pf, previous_pf, time_period)
    pf_type = "지상" if time_period == "daytime" else "진상"
    time_range = "(09-23시)" if time_period == "daytime" else "(23-09시)"
    
    card_html = f"""
    <div class="time-period-card {card_class}">
        <div class="card-title">{icon} {period_name} 역률 {time_range}</div>
        <div class="card-value"><span class="traffic-light">{traffic_light}</span>{pf_type} {current_pf:.1f}%</div>
        <div class="card-change">{message}</div>
        <div class="card-period">한전 요금체계 기준</div>
    </div>
    """
    return card_html

def create_summary_table(current_data, period_type="일"):
    """요약 테이블 생성"""
    numeric_columns = [("전력사용량(kWh)", "kWh"), ("지상무효전력량(kVarh)", "kVarh"), ("진상무효전력량(kVarh)", "kVarh"),
                      ("탄소배출량(tCO2)", "tCO2"), ("지상역률(%)", "%"), ("진상역률(%)", "%"), ("전기요금(원)", "원")]
    ratio_cols = {"지상역률(%)", "진상역률(%)"}

    rows = []
    for col, unit in numeric_columns:
        if col in ratio_cols:
            val = current_data[col].mean()
        else:
            val = current_data[col].sum()
        name = col.split("(")[0]
        rows.append({"항목": name, f"현재{period_type} 값": f"{val:.2f}", "단위": unit})
    return pd.DataFrame(rows)

def create_comparison_table(current_data, previous_data, period_type="일"):
    """비교 테이블 생성"""
    comparison_dict = {"항목": [], f"현재{period_type}": [], f"이전{period_type}": [], "변화량": [], "변화율(%)": []}
    numeric_columns = ["전력사용량(kWh)", "지상무효전력량(kVarh)", "진상무효전력량(kVarh)", "탄소배출량(tCO2)", "지상역률(%)", "진상역률(%)", "전기요금(원)"]

    for col in numeric_columns:
        if col in ["지상역률(%)", "진상역률(%)"]:
            current_val = current_data[col].mean()
            previous_val = previous_data[col].mean()
        else:
            current_val = current_data[col].sum()
            previous_val = previous_data[col].sum()

        change = current_val - previous_val
        change_pct = (change / previous_val * 100) if previous_val != 0 else 0

        comparison_dict["항목"].append(col)
        comparison_dict[f"현재{period_type}"].append(f"{current_val:.2f}")
        comparison_dict[f"이전{period_type}"].append(f"{previous_val:.2f}")
        comparison_dict["변화량"].append(f"{change:+.2f}")
        comparison_dict["변화율(%)"].append(f"{change_pct:+.1f}%")
    return pd.DataFrame(comparison_dict)

# ========== 4. 개선된 보고서 생성 함수 ==========
def create_comprehensive_docx_report_with_charts(df, current_data, daily_data, selected_date, view_type="월별", selected_month=1, period_label="전체"):
    """현재 화면 설정에 따른 동적 보고서 생성"""
    doc = Document()
    
    # 전체 문서에 테두리 추가
    sections = doc.sections
    for section in sections:
        sectPr = section._sectPr
        pgBorders = sectPr.xpath('.//w:pgBorders')
        if not pgBorders:
            from docx.oxml import parse_xml
            pgBorders_xml = '''
            <w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                         w:offsetFrom="page">
                <w:top w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:left w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:bottom w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:right w:val="single" w:sz="12" w:space="24" w:color="auto"/>
            </w:pgBorders>
            '''
            pgBorders = parse_xml(pgBorders_xml)
            sectPr.append(pgBorders)
    
    # 보고서 헤더 테이블 (6행으로 축소)
    header_table = doc.add_table(rows=6, cols=4)
    header_table.style = 'Table Grid'
    
    # 제목 행
    title_cell = header_table.rows[0].cells[0]
    title_cell.merge(header_table.rows[0].cells[3])
    title_cell.text = "보 고 서"
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para = title_cell.paragraphs[0]
    title_para.runs[0].font.size = Cm(0.8)
    title_para.runs[0].bold = True
    
    # 헤더 정보 입력
    header_data = [
        ("보고처", "에너지관리팀", "보고서명", f"전력 분석 보고서 ({period_label})"),
        ("장소", "본사", "취급분류", "○기밀 ●보통"),
        ("작성일자", datetime.now().strftime("%Y년 %m월 %d일"), "작성자", "홍Zoo형"),
        ("참가자", "에너지관리팀, 시설관리팀, 경영진", "", ""),
        ("자료출처", "전력량계 실시간 데이터, 한국전력공사 요금체계", "", "")
    ]
    
    for i, (col1, val1, col2, val2) in enumerate(header_data, 1):
        header_table.rows[i].cells[0].text = col1
        header_table.rows[i].cells[1].text = val1
        if col2:
            header_table.rows[i].cells[2].text = col2
            header_table.rows[i].cells[3].text = val2
        else:
            header_table.rows[i].cells[1].merge(header_table.rows[i].cells[3])
    
    doc.add_paragraph()
    
    # 보고내용 제목
    content_title = doc.add_heading('보고내용', level=1)
    content_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # === 1. 기간별 분석 (화면 설정에 따라 동적) ===
    if view_type == "월별":
        doc.add_heading(f'1. {selected_month}월 전력 사용 분석', level=2)
        
        if not current_data.empty:
            total_kwh = current_data["전력사용량(kWh)"].sum()
            total_cost = current_data["전기요금(원)"].sum()
            avg_pf = current_data["지상역률(%)"].mean()
            total_carbon = current_data["탄소배출량(tCO2)"].sum()
            avg_price = total_cost / total_kwh if total_kwh > 0 else 0
            
            doc.add_paragraph(f"□ {selected_month}월 총 전력사용량: {total_kwh:,.1f} kWh")
            doc.add_paragraph(f"□ {selected_month}월 총 전기요금: {total_cost:,.0f} 원")
            doc.add_paragraph(f"□ {selected_month}월 평균 단가: {avg_price:.1f} 원/kWh")
            doc.add_paragraph(f"□ {selected_month}월 평균 역률: {avg_pf:.1f}%")
            doc.add_paragraph(f"□ {selected_month}월 탄소배출량: {total_carbon:.2f} tCO2")
            
            # 월별 트렌드 차트 생성
            try:
                monthly_summary = df.groupby('년월').agg({
                    '전력사용량(kWh)': 'sum',
                    '전기요금(원)': 'sum'
                }).reset_index()
                monthly_summary['년월_str'] = monthly_summary['년월'].astype(str)
                chart_data = monthly_summary[['년월_str', '전력사용량(kWh)']].tail(6)
                
                chart_img = create_matplotlib_chart(
                    chart_data, 
                    chart_type="line", 
                    title="월별 전력사용량 추이 (최근 6개월)",
                    xlabel="월",
                    ylabel="전력사용량 (kWh)",
                    figsize=(10, 5)
                )
                
                doc.add_paragraph()
                doc.add_paragraph("【월별 전력사용량 추이 그래프】")
                doc.add_picture(chart_img, width=Inches(6))
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f"※ 월별 차트 생성 실패: {str(e)}")
    
    else:  # 일별 분석
        doc.add_heading(f'1. {period_label} 전력 사용 분석', level=2)
        
        if not current_data.empty:
            total_kwh = current_data["전력사용량(kWh)"].sum()
            total_cost = current_data["전기요금(원)"].sum()
            avg_pf = current_data["지상역률(%)"].mean()
            total_carbon = current_data["탄소배출량(tCO2)"].sum()
            avg_price = total_cost / total_kwh if total_kwh > 0 else 0
            
            doc.add_paragraph(f"□ 기간 총 전력사용량: {total_kwh:,.1f} kWh")
            doc.add_paragraph(f"□ 기간 총 전기요금: {total_cost:,.0f} 원")
            doc.add_paragraph(f"□ 기간 평균 단가: {avg_price:.1f} 원/kWh")
            doc.add_paragraph(f"□ 기간 평균 역률: {avg_pf:.1f}%")
            doc.add_paragraph(f"□ 기간 탄소배출량: {total_carbon:.2f} tCO2")
    
    # === 2. 특정일 시간별 분석 ===
    if daily_data is not None and not daily_data.empty:
        doc.add_heading(f'2. {selected_date} 시간별 분석', level=2)
        
                
        # 최대 사용 시간 정보
        hourly_summary = daily_data.groupby('시간').agg({
            '전력사용량(kWh)': 'sum',
            '전기요금(원)': 'sum'
        }).reset_index()
        peak_hour = hourly_summary.loc[hourly_summary['전력사용량(kWh)'].idxmax()]
        doc.add_paragraph(f"□ 최대 사용시간: {int(peak_hour['시간'])}시 ({peak_hour['전력사용량(kWh)']:.1f} kWh)")
        doc.add_paragraph(f"□ 해당 시간 전기요금: {peak_hour['전기요금(원)']:,.0f} 원")
    
    # === 3. 전일 대비 역률 요금 분석 (텍스트) ===
    doc.add_heading('3. 전일 대비 역률 요금 분석', level=2)
    
    if daily_data is not None and not daily_data.empty:
        # 시간대별 역률 분석
        daytime_data = daily_data[(daily_data['시간'] >= 9) & (daily_data['시간'] < 23)]
        nighttime_data = daily_data[(daily_data['시간'] >= 23) | (daily_data['시간'] < 9)]
        
        if len(daytime_data) > 0:
            daytime_pf = daytime_data['지상역률(%)'].mean()
            doc.add_paragraph(f"□ 주간 평균 지상역률 (09-23시): {daytime_pf:.1f}%")
            
            adjusted_pf = max(60, min(95, daytime_pf))
            if adjusted_pf >= 90:
                rate_impact = -(adjusted_pf - 90) * 0.5
                impact_text = f"감액 {abs(rate_impact):.1f}%"
            else:
                rate_impact = (90 - adjusted_pf) * 0.5
                impact_text = f"추가요금 {rate_impact:.1f}%"
            doc.add_paragraph(f"  - 한전 요금 영향: {impact_text}")
        
        if len(nighttime_data) > 0:
            nighttime_pf = nighttime_data['진상역률(%)'].mean()
            if nighttime_pf > 0:
                doc.add_paragraph(f"□ 야간 평균 진상역률 (23-09시): {nighttime_pf:.1f}%")
                adjusted_pf = max(60, nighttime_pf)
                if adjusted_pf >= 95:
                    impact_text = "추가요금 없음"
                else:
                    rate_impact = (95 - adjusted_pf) * 0.5
                    impact_text = f"추가요금 {rate_impact:.1f}%"
                doc.add_paragraph(f"  - 한전 요금 영향: {impact_text}")
            else:
                doc.add_paragraph(f"□ 야간 지상역률 운전 (23-09시): 정상 운전")
                doc.add_paragraph(f"  - 한전 요금 영향: 추가요금 없음")
    
    # === 4. 상세 비교 데이터 (표) ===
    doc.add_heading('4. 상세 비교 데이터', level=2)
    
    if daily_data is not None and not daily_data.empty:
        # 시간별 상세 테이블
        doc.add_paragraph("【시간별 상세 현황표】")
        table_hourly = doc.add_table(rows=1, cols=4)
        table_hourly.style = 'Table Grid'
        hdr_cells = table_hourly.rows[0].cells
        hdr_cells[0].text = '시간'
        hdr_cells[1].text = '전력사용량(kWh)'
        hdr_cells[2].text = '전기요금(원)'
        hdr_cells[3].text = '지상역률(%)'
        
        hourly_summary = daily_data.groupby('시간').agg({
            '전력사용량(kWh)': 'sum',
            '전기요금(원)': 'sum',
            '지상역률(%)': 'mean'
        }).round(2)
        
        # 상위 12시간만 표시
        top_hours = hourly_summary.sort_values('전력사용량(kWh)', ascending=False).head(12)
        for hour in top_hours.index:
            row_cells = table_hourly.add_row().cells
            row_cells[0].text = f"{hour}시"
            row_cells[1].text = f"{top_hours.loc[hour, '전력사용량(kWh)']:,.1f}"
            row_cells[2].text = f"{top_hours.loc[hour, '전기요금(원)']:,.0f}"
            row_cells[3].text = f"{top_hours.loc[hour, '지상역률(%)']:.1f}"

    # === 5. 시간대별 작업유형별 전기요금 현황 (차트) ===
    doc.add_heading('5. 시간대별 작업유형별 전기요금 현황', level=2)
    
    analysis_data = daily_data if (daily_data is not None and not daily_data.empty) else df
    
    # 막대그래프 생성 (시간대별 작업유형별 전기요금)
    try:
        hourly_worktype = analysis_data.groupby(['시간', '작업유형'])['전기요금(원)'].sum().unstack(fill_value=0)
        
        # matplotlib로 스택 바 차트 생성
        fig, ax = plt.subplots(figsize=(12, 6))
        colors = {'Light_Load': '#4CAF50', 'Medium_Load': '#FF9800', 'Maximum_Load': '#F44336'}
        names = {'Light_Load': '경부하', 'Medium_Load': '중간부하', 'Maximum_Load': '최대부하'}
        
        bottom = np.zeros(len(hourly_worktype.index))
        for work_type in ['Light_Load', 'Medium_Load', 'Maximum_Load']:
            if work_type in hourly_worktype.columns:
                ax.bar(hourly_worktype.index, hourly_worktype[work_type], 
                      bottom=bottom, label=names[work_type], 
                      color=colors[work_type])
                bottom += hourly_worktype[work_type]
        
        ax.set_title('시간대별 작업유형별 전기요금 현황', fontsize=14, fontweight='bold')
        ax.set_xlabel('시간', fontsize=12)
        ax.set_ylabel('전기요금 (원)', fontsize=12)
        ax.legend()
        ax.grid(True, alpha=0.3)
        plt.tight_layout()
        
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight', 
                    facecolor='white', edgecolor='none')
        img_buffer.seek(0)
        plt.close()
        
        doc.add_paragraph("【시간대별 작업유형별 전기요금 막대그래프】")
        doc.add_picture(img_buffer, width=Inches(6.5))
        doc.add_paragraph()
    except Exception as e:
        doc.add_paragraph(f"※ 막대그래프 생성 실패: {str(e)}")
    
    # 파이차트 생성 (작업유형별 전력사용량)
    try:
        worktype_stats = analysis_data.groupby('작업유형')['전력사용량(kWh)'].sum().reset_index()
        worktype_stats['작업유형_한글'] = worktype_stats['작업유형'].map({
            'Light_Load': '경부하',
            'Medium_Load': '중간부하', 
            'Maximum_Load': '최대부하'
        })
        pie_data = worktype_stats[['작업유형_한글', '전력사용량(kWh)']]
        
        chart_img = create_matplotlib_chart(
            pie_data,
            chart_type="pie",
            title="작업유형별 전력사용량 분포",
            figsize=(8, 6)
        )
        
        doc.add_paragraph("【작업유형별 전력사용량 파이차트】")
        doc.add_picture(chart_img, width=Inches(5))
        doc.add_paragraph()
    except Exception as e:
        doc.add_paragraph(f"※ 파이차트 생성 실패: {str(e)}")
    
    # === 6. 작업유형별 상세 분석 (표) ===
    doc.add_heading('6. 작업유형별 상세 분석', level=2)
    
    worktype_detailed = analysis_data.groupby('작업유형').agg({
        '전력사용량(kWh)': 'sum',
        '전기요금(원)': 'sum',
        '지상역률(%)': 'mean',
        '탄소배출량(tCO2)': 'sum'
    }).round(2)
    
    doc.add_paragraph("【작업유형별 상세 현황표】")
    table_worktype = doc.add_table(rows=1, cols=5)
    table_worktype.style = 'Table Grid'
    hdr_cells = table_worktype.rows[0].cells
    hdr_cells[0].text = '작업유형'
    hdr_cells[1].text = '전력사용량(kWh)'
    hdr_cells[2].text = '전기요금(원)'
    hdr_cells[3].text = '평균역률(%)'
    hdr_cells[4].text = '탄소배출량(tCO2)'
    
    for work_type in worktype_detailed.index:
        row_cells = table_worktype.add_row().cells
        type_map = {'Light_Load': '경부하', 'Medium_Load': '중간부하', 'Maximum_Load': '최대부하'}
        row_cells[0].text = type_map.get(work_type, work_type)
        row_cells[1].text = f"{worktype_detailed.loc[work_type, '전력사용량(kWh)']:,.1f}"
        row_cells[2].text = f"{worktype_detailed.loc[work_type, '전기요금(원)']:,.0f}"
        row_cells[3].text = f"{worktype_detailed.loc[work_type, '지상역률(%)']:.1f}"
        row_cells[4].text = f"{worktype_detailed.loc[work_type, '탄소배출량(tCO2)']:.2f}"
    
    # === 부록: 용어 설명 ===
    doc.add_page_break()
    doc.add_heading('부록 - 전력 관련 용어 설명', level=1)
    
    doc.add_paragraph("【전력 측정 지표 설명】")
    
    doc.add_paragraph("□ 전력사용량(kWh)")
    doc.add_paragraph("  - 실제 소비한 전력량으로 전기요금 산정의 기본 단위")
    doc.add_paragraph("  - 1kWh = 1000W의 전력을 1시간 사용한 양")
    
    doc.add_paragraph("□ 지상무효전력량(kVarh)")
    doc.add_paragraph("  - 유도성 부하(모터, 변압기 등)에서 발생하는 무효전력")
    doc.add_paragraph("  - 전압이 전류보다 앞서는 경우로, 역률 저하의 주요 원인")
    
    doc.add_paragraph("□ 진상무효전력량(kVarh)")
    doc.add_paragraph("  - 용량성 부하(콘덴서 등)에서 발생하는 무효전력")
    doc.add_paragraph("  - 전류가 전압보다 앞서는 경우")
    
    doc.add_paragraph("□ 지상역률(%)")
    doc.add_paragraph("  - 유효전력과 피상전력의 비율로 전력 효율을 나타냄")
    doc.add_paragraph("  - 한전 기준: 90% 이상 유지 시 요금 감액, 미만 시 할증")
    
    doc.add_paragraph("□ 진상역률(%)")
    doc.add_paragraph("  - 진상 운전 시의 역률로 주로 야간 경부하 시간대에 발생")
    doc.add_paragraph("  - 한전 기준: 95% 이상 유지 시 추가요금 없음")
    
    doc.add_paragraph("□ 작업유형 구분")
    doc.add_paragraph("  - 경부하(Light_Load): 23:00~09:00, 전력사용량이 적은 시간대")
    doc.add_paragraph("  - 중간부하(Medium_Load): 09:00~10:00, 12:00~13:00, 17:00~23:00")
    doc.add_paragraph("  - 최대부하(Maximum_Load): 10:00~12:00, 13:00~17:00, 요금이 가장 높은 시간대")
    
    doc.add_paragraph("□ 탄소배출량(tCO2)")
    doc.add_paragraph("  - 전력 사용으로 인한 이산화탄소 배출량")
    doc.add_paragraph("  - 환경부 고시 배출계수 적용하여 산정")
    
    # === 마무리 ===
    doc.add_paragraph()
    doc.add_heading('보고서 결론', level=1)
    
    total_kwh = df["전력사용량(kWh)"].sum()
    total_cost = df["전기요금(원)"].sum()
    avg_pf = df["지상역률(%)"].mean()
    
    doc.add_paragraph("【종합 분석 결과】")
    doc.add_paragraph(f"□ 전체 분석기간 전력사용량: {total_kwh:,.1f} kWh")
    doc.add_paragraph(f"□ 전체 분석기간 전기요금: {total_cost:,.0f} 원")
    doc.add_paragraph(f"□ 평균 역률: {avg_pf:.1f}%")
    
    if avg_pf < 90:
        doc.add_paragraph(f"□ 역률 개선 필요: 현재 {avg_pf:.1f}%로 90% 미만")
        monthly_avg_cost = total_cost / len(df['년월'].unique()) if len(df['년월'].unique()) > 0 else total_cost
        doc.add_paragraph(f"□ 역률 개선을 통한 예상 절약효과: 월 약 {monthly_avg_cost * 0.05:,.0f} 원")
    else:
        doc.add_paragraph(f"□ 역률 상태 양호: 현재 {avg_pf:.1f}%로 기준치 이상 유지")
    
    doc.add_paragraph()
    doc.add_paragraph("【향후 계획】")
    doc.add_paragraph("□ 지속적인 전력 사용 패턴 모니터링")
    doc.add_paragraph("□ 역률 개선을 통한 전기요금 절감 방안 검토")
    doc.add_paragraph("□ 부하 분산을 통한 최대부하 시간대 사용량 최적화")
    doc.add_paragraph("□ 월별 정기 분석을 통한 에너지 효율 관리 강화")
    
    doc.add_paragraph()
    doc.add_paragraph("이상으로 전력 분석 보고를 마치겠습니다.")
    
    return doc

# ========== 5. 메인 함수 (원래 코드 그대로 유지) ==========
def main():
    st.title("과거 전기요금 분석 보고서")

    st.markdown('<div class="header-style"><h1>통합 전력 분석</h1></div>', unsafe_allow_html=True)

    df = load_data()
    if df is None:
        st.stop()

    st.sidebar.header("분석 설정")
    filtered_df = df.copy()
    date_range = (df["날짜"].min(), df["날짜"].max())
    work_types = df["작업유형"].unique()
    
    st.sidebar.subheader("상세 분석 옵션")
    numeric_columns = ["전력사용량(kWh)", "지상무효전력량(kVarh)", "진상무효전력량(kVarh)", "탄소배출량(tCO2)", "지상역률(%)", "진상역률(%)", "전기요금(원)"]
    col1_select = st.sidebar.selectbox("첫 번째 분석 컬럼", numeric_columns, index=0)
    col2_select = st.sidebar.selectbox("두 번째 분석 컬럼", numeric_columns, index=6)

    # 보고서 생성 옵션
    st.sidebar.subheader("보고서 생성 옵션")
        
    if st.sidebar.button("보고서 생성", key="generate_complete_report"):
        with st.sidebar:
            with st.spinner("보고서 생성 중..."):
                try:
                    # 현재 설정된 분석 조건 가져오기
                    view_type = st.session_state.get('analysis_period', '월별')
                    
                    if view_type == "월별":
                        selected_month = st.session_state.get('month_selector', 1)
                        current_year = filtered_df["년월"].dt.year.max()
                        current_data = filtered_df[
                            (filtered_df["년월"].dt.year == current_year) &
                            (filtered_df["년월"].dt.month == selected_month)
                        ]
                        period_label = f"{selected_month}월"
                    else:
                        # 일별 분석의 경우
                        selected_range = st.session_state.get('period_range_selector', None)
                        if selected_range and len(selected_range) == 2:
                            start_day, end_day = selected_range
                            current_data = filtered_df[
                                (filtered_df["날짜"] >= start_day) & 
                                (filtered_df["날짜"] <= end_day)
                            ]
                            period_label = f"{start_day} ~ {end_day} 기간"
                        else:
                            current_data = filtered_df
                            period_label = "전체 기간"
                    
                    # 최근 날짜 데이터
                    latest_date = filtered_df["날짜"].max()
                    daily_data = filtered_df[filtered_df["날짜"] == latest_date]
                    
                    # 보고서 생성
                    doc = create_comprehensive_docx_report_with_charts(
                        filtered_df, current_data, daily_data, latest_date, 
                        view_type, selected_month if view_type == "월별" else None, period_label
                    )
                    
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"전기요금_분석보고서_{timestamp}.docx"
                    
                    st.download_button(
                        label="보고서 다운로드",
                        data=doc_buffer.getvalue(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_complete_report"
                    )
                    
                    st.success("보고서 생성 완료!")
                    st.info(f"파일명: {filename}")
                    st.info("분석 보고서가 생성되었습니다!")
                    
                except Exception as e:
                    st.error(f"보고서 생성 중 오류 발생: {str(e)}")
                    st.info("오류가 지속되면 다른 날짜나 기간을 선택해보세요.")

    summary_data = filtered_df.copy()
    period_label = "전체"

    # 필터링 옵션
    filter_col1, filter_col2, filter_col3, filter_col4 = st.columns([0.8, 0.5, 1, 0.7])
    
    with filter_col1:
        view_type = st.selectbox("분석 기간", ["월별", "일별"], key="analysis_period")
    
    with filter_col2:
        if view_type == "월별":
            current_year = filtered_df["년월"].dt.year.max()
            months = list(range(1, 13))
            default_month = filtered_df["년월"].dt.month.max()
            selected_month = st.selectbox("월", months, index=int(default_month) - 1, key="month_selector")
        else:
            st.markdown("")
    
    with filter_col3:
        if view_type == "일별":
            from datetime import date
            selected_range = st.date_input("기간", value=(date(2024, 1, 1), date(2024, 1, 5)),
                                         min_value=date_range[0] if len(date_range) == 2 else df["날짜"].min(),
                                         max_value=date_range[1] if len(date_range) == 2 else df["날짜"].max(),
                                         key="period_range_selector")
        else:
            st.markdown("")
    
    with filter_col4:
        st.markdown("")
    
    current_data = pd.DataFrame()
    previous_data = pd.DataFrame()

    # 데이터 처리 로직
    if view_type == "월별":
        current_data = filtered_df[(filtered_df["년월"].dt.year == current_year) & (filtered_df["년월"].dt.month == selected_month)]
        summary_data = current_data
        period_label = f"{selected_month}월"

        if selected_month > 1:
            prev_year = current_year
            prev_month = selected_month - 1
        else:
            prev_year = current_year - 1
            prev_month = 12
        previous_data = filtered_df[(filtered_df["년월"].dt.year == prev_year) & (filtered_df["년월"].dt.month == prev_month)]

    else:
        if not isinstance(selected_range, tuple) or len(selected_range) != 2:
            st.warning("날짜 범위를 선택해주세요")
        else:
            start_day, end_day = selected_range
            if start_day > end_day:
                st.warning("시작 날짜가 종료 날짜보다 이후입니다.")
            else:
                period_df = filtered_df[(filtered_df["날짜"] >= start_day) & (filtered_df["날짜"] <= end_day)]
                if period_df.empty:
                    st.info(f"{start_day} ~ {end_day} 구간에는 데이터가 없습니다.")
                else:
                    current_data = period_df
                    summary_data = period_df
                    period_label = f"{start_day} ~ {end_day} 기간"
                    
                    days = (end_day - start_day).days + 1
                    prev_start = start_day - timedelta(days=days)
                    prev_end = start_day - timedelta(days=1)
                    previous_data = filtered_df[(filtered_df["날짜"] >= prev_start) & (filtered_df["날짜"] <= prev_end)]

    # 주요 지표 카드
    if not summary_data.empty:
        main_metrics_card = create_main_metrics_card(summary_data, period_label)
        st.markdown(main_metrics_card, unsafe_allow_html=True)

    # 차트 섹션
    if view_type == "월별":
        monthly_data = (filtered_df.groupby("년월").agg({
            col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
            col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
        }).reset_index())
        monthly_data["년월_str"] = monthly_data["년월"].astype(str)

        fig = create_dual_axis_chart(monthly_data, "년월_str", col1_select, col2_select,
                                   f"월별 {col1_select} vs {col2_select} 비교", "월", col1_select, col2_select)
        st.plotly_chart(fig, use_container_width=True)

    else:
        if isinstance(current_data, pd.DataFrame) and not current_data.empty:
            daily_data = (period_df.groupby("날짜").agg({
                col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
                col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
            }).reset_index())

            fig = create_dual_axis_chart(daily_data, "날짜", col1_select, col2_select,
                                       f"{start_day} ~ {end_day} 날짜별 {col1_select} vs {col2_select}",
                                       "날짜", col1_select, col2_select)
            st.plotly_chart(fig, use_container_width=True)

    # 월별 분석일 때 비교 테이블
    if view_type == "월별" and not current_data.empty and not previous_data.empty:
        st.subheader("전월 대비 분석")
        comparison_df = create_comparison_table(current_data, previous_data, "월")
        st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
        st.dataframe(comparison_df, use_container_width=True, hide_index=True)
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")

    # 특정일 시간별 분석
    st.markdown('<div class="section-header"><h2>특정일 시간별 에너지 사용 분석</h2></div>', unsafe_allow_html=True)

    col1, col2 = st.columns([3, 1])
    daily_df = pd.DataFrame()

    with col1:
        available_dates = sorted(filtered_df["날짜"].unique())
        if available_dates:
            min_d, max_d = available_dates[0], available_dates[-1]
            default_d = max_d
            selected_date = st.date_input("분석할 날짜 선택", value=default_d, min_value=min_d, max_value=max_d, key="daily_date_selector")

            daily_df = filtered_df[filtered_df["날짜"] == selected_date]
            if daily_df.empty:
                st.warning(f"{selected_date} 데이터가 없습니다.")
            else:
                hourly_data = (daily_df.groupby("시간").agg({
                    col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
                    col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
                }).reset_index())

                full_hours = pd.DataFrame({"시간": list(range(24))})
                hourly_data = pd.merge(full_hours, hourly_data, on="시간", how="left").fillna(0)

                fig = create_dual_axis_chart(hourly_data, "시간", col1_select, col2_select,
                                           f"{selected_date} 시간별 {col1_select} vs {col2_select} 비교",
                                           "시간", col1_select, col2_select, add_time_zones=True)

                fig.update_xaxes(tickmode="linear", tick0=0, dtick=1, title_text="시간")
                st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("선택된 조건에 맞는 데이터가 없습니다.")

    with col2:
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.subheader("전일 대비 역률 요금")

        if available_dates and selected_date in available_dates:
            try:
                date_idx = available_dates.index(selected_date)
                if date_idx > 0:
                    previous_date = available_dates[date_idx - 1]
                    previous_daily_df = filtered_df[filtered_df["날짜"] == previous_date]

                    if not daily_df.empty and not previous_daily_df.empty:
                        current_daytime = daily_df[(daily_df['시간'] >= 9) & (daily_df['시간'] < 23)]
                        previous_daytime = previous_daily_df[(previous_daily_df['시간'] >= 9) & (previous_daily_df['시간'] < 23)]
                        current_nighttime = daily_df[(daily_df['시간'] >= 23) | (daily_df['시간'] < 9)]
                        previous_nighttime = previous_daily_df[(previous_daily_df['시간'] >= 23) | (previous_daily_df['시간'] < 9)]
                        
                        if len(current_daytime) > 0:
                            current_daytime_raw = current_daytime['지상역률(%)'].mean()
                            current_daytime_pf = max(60, min(95, current_daytime_raw))
                        else:
                            current_daytime_pf = 90
                        
                        if len(previous_daytime) > 0:
                            previous_daytime_raw = previous_daytime['지상역률(%)'].mean()
                            previous_daytime_pf = max(60, min(95, previous_daytime_raw))
                        else:
                            previous_daytime_pf = 90
                        
                        if len(current_nighttime) > 0:
                            current_leading_raw = current_nighttime['진상역률(%)'].mean()
                            if current_leading_raw > 0:
                                current_nighttime_pf = max(60, current_leading_raw)
                            else:
                                current_nighttime_pf = 100
                        else:
                            current_nighttime_pf = 100
                        
                        if len(previous_nighttime) > 0:
                            previous_leading_raw = previous_nighttime['진상역률(%)'].mean()
                            if previous_leading_raw > 0:
                                previous_nighttime_pf = max(60, previous_leading_raw)
                            else:
                                previous_nighttime_pf = 100
                        else:
                            previous_nighttime_pf = 100
                        
                        daytime_card = create_simple_power_factor_card("주간", "", current_daytime_pf, previous_daytime_pf, "daytime", "daytime-card")
                        nighttime_card = create_simple_power_factor_card("야간", "", current_nighttime_pf, previous_nighttime_pf, "nighttime", "nighttime-card")
                        
                        st.markdown(daytime_card, unsafe_allow_html=True)
                        st.markdown(nighttime_card, unsafe_allow_html=True)
                    else:
                        st.info("선택된 날짜 또는 전일 데이터가 없습니다.")
                else:
                    if not daily_df.empty:
                        summary_df = create_summary_table(daily_df, "일")
                        st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
                        st.dataframe(summary_df, use_container_width=True, hide_index=True)
                        st.markdown("</div>", unsafe_allow_html=True)
                        st.info("첫 번째 날짜로 전일 데이터가 없어 비교할 수 없습니다.")
                    else:
                        st.info("선택된 날짜의 데이터가 없습니다.")
            except (ValueError, IndexError):
                st.info("이전 날짜를 찾을 수 없습니다.")

    # 상세 비교 데이터 표
    if available_dates and selected_date in available_dates:
        try:
            date_idx = available_dates.index(selected_date)
            if date_idx > 0:
                previous_date = available_dates[date_idx - 1]
                previous_daily_df = filtered_df[filtered_df["날짜"] == previous_date]
                
                if not daily_df.empty and not previous_daily_df.empty:
                    st.subheader("상세 비교 데이터")
                    comparison_df = create_comparison_table(daily_df, previous_daily_df, "일")
                    st.dataframe(comparison_df, use_container_width=True, hide_index=True)
        except (ValueError, IndexError):
            pass

    st.markdown("---")

    # 시간대별 현황 차트
    if not daily_df.empty:
        col_chart1, col_chart2 = st.columns([2, 1])
        with col_chart1:
            st.plotly_chart(create_hourly_stack_chart(daily_df), use_container_width=True)
        with col_chart2:
            st.plotly_chart(create_concentric_donut_chart(daily_df), use_container_width=True)

        st.subheader(f"{selected_date} 작업유형별 상세 분석")
        worktype_stats = (daily_df.groupby("작업유형").agg(
            전력사용량_합계=("전력사용량(kWh)", "sum"),
            전기요금_합계=("전기요금(원)", "sum"),
            평균_지상역률=("지상역률(%)", "mean"),
            탄소배출량_합계=("탄소배출량(tCO2)", "sum")
        ).round(2))
        st.dataframe(worktype_stats, use_container_width=True)
    else:
        col_chart1, col_chart2 = st.columns([2, 1])
        with col_chart1:
            st.plotly_chart(create_hourly_stack_chart(filtered_df), use_container_width=True)
        with col_chart2:
            st.plotly_chart(create_concentric_donut_chart(filtered_df), use_container_width=True)

        st.subheader("전체 작업유형별 상세 분석")
        worktype_stats = (filtered_df.groupby("작업유형").agg(
            전력사용량_합계=("전력사용량(kWh)", "sum"),
            전기요금_합계=("전기요금(원)", "sum"),
            평균_지상역률=("지상역률(%)", "mean"),
            탄소배출량_합계=("탄소배출량(tCO2)", "sum")
        ).round(2))
        st.dataframe(worktype_stats, use_container_width=True)

    st.markdown("---")

if __name__ == "__main__":
    main()