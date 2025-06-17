def create_simple_power_factor_card(period_name, icon, current_pf, previous_pf, time_period, card_class):
    """역률 카드 생성"""
    traffic_light, message, style_class = get_traffic_light_and_message(current_pf, previous_pf, time_period)
    pf_type = "지상" if time_period == "daytime" else "진상"
    time_range = "(09-23시)" if time_period == "daytime" else "(23-09시)"
    
    # 역률 상태에 따른 색상 결정 (수정됨)
    if time_period == "daytime":
        # 주간 지상역률 기준
        if current_pf >= 90:
            status_color = "#34a853"
            status_text = "우수"
        elif current_pf >= 85:
            status_color = "#fbbc04"
            status_text = "양호"
        else:
            status_color = "#ea4335"
            status_text = "개선필요"
    else:
        # 야간 진상역률 기준
        if current_pf >= 95:
            status_color = "#34a853"
            status_text = "우수"
        elif current_pf >= 85:
            status_color = "#fbbc04"
            status_text = "양호"
        else:
            status_color = "#ea4335"
            status_text = "개선필요"
    
    card_html = f"""
    <div class="time-period-card {card_class}">
        <div class="card-title">{icon} {period_name} 역률 관리 {time_range}</div>
        <div class="card-value">
            <span class="traffic-light">{traffic_light}</span>
            {pf_type} {current_pf:.1f}%
        </div>
        <div class="card-change" style="{style_class}">{message}</div>
        <div style="margin-top: 8px; padding: 4px 8px; border-radius: 4px; display: inline-block; 
             background: {status_color}; color: white; font-size: 0.8rem; font-weight: 600;">
            상태: {status_text}
        </div>
        <div class="card-period">한국전력공사 요금체계 기준</div>
    </div>
    """
    return card_html

def create_summary_table(current_data, period_type="일"):
    """요약 테이블 생성"""
    numeric_columns = [
        ("전력사용량(kWh)", "kWh"), ("지상무효전력량(kVarh)", "kVarh"), 
        ("진상무효전력량(kVarh)", "kVarh"), ("탄소배출량(tCO2)", "tCO₂"), 
        ("지상역률(%)", "%"), ("진상역률(%)", "%"), ("전기요금(원)", "원")
    ]
    ratio_cols = {"지상역률(%)", "진상역률(%)"}

    rows = []
    for col, unit in numeric_columns:
        if col in current_data.columns:
            if col in ratio_cols:
                val = current_data[col].mean()
            else:
                val = current_data[col].sum()
            name = col.split("(")[0]
            rows.append({"항목": name, f"현재{period_type} 값": f"{val:.2f}", "단위": unit})
    return pd.DataFrame(rows)

def create_comparison_table(current_data, previous_data, period_type="일"):
    """비교 테이블 생성"""
    comparison_dict = {
        "항목": [], f"현재{period_type}": [], f"이전{period_type}": [], 
        "변화량": [], "변화율(%)": [], "상태": []
    }
    numeric_columns = [
        "전력사용량(kWh)", "지상무효전력량(kVarh)", "진상무효전력량(kVarh)", 
        "탄소배출량(tCO2)", "지상역률(%)", "진상역률(%)", "전기요금(원)"
    ]

    for col in numeric_columns:
        if col in current_data.columns and col in previous_data.columns:
            if col in ["지상역률(%)", "진상역률(%)"]:
                current_val = current_data[col].mean()
                previous_val = previous_data[col].mean()
            else:
                current_val = current_data[col].sum()
                previous_val = previous_data[col].sum()

            change = current_val - previous_val
            change_pct = (change / previous_val * 100) if previous_val != 0 else 0
            
            # 상태 결정
            if col == "전기요금(원)":
                status = "🔴 증가" if change > 0 else "🟢 감소" if change < 0 else "🟡 동일"
            elif col in ["지상역률(%)", "진상역률(%)"]:
                status = "🟢 개선" if change > 0 else "🔴 악화" if change < 0 else "🟡 동일"
            else:
                status = "🔴 증가" if change > 0 else "🟢 감소" if change < 0 else "🟡 동일"

            comparison_dict["항목"].append(col.replace("(", " ("))
            comparison_dict[f"현재{period_type}"].append(f"{current_val:.2f}")
            comparison_dict[f"이전{period_type}"].append(f"{previous_val:.2f}")
            comparison_dict["변화량"].append(f"{change:+.2f}")
            comparison_dict["변화율(%)"].append(f"{change_pct:+.1f}%")
            comparison_dict["상태"].append(status)
        
    return pd.DataFrame(comparison_dict)

def create_comprehensive_docx_report_with_charts(df, current_data, daily_data, selected_date, view_type="월별", selected_month=1, period_label="전체"):
    """현재 화면 설정에 따른 동적 보고서 생성"""
    doc = Document()
    
    # 전체 문서에 테두리 추가
    sections = doc.sections
    for section in sections:
        sectPr = section._sectPr
        pgBorders = sectPr.xpath('.//w:pgBorders')
        if not pgBorders:
            from docx.oxml import parse_xml
            pgBorders_xml = '''
            <w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                         w:offsetFrom="page">
                <w:top w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:left w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:bottom w:val="single" w:sz="12" w:space="24" w:color="auto"/>
                <w:right w:val="single" w:sz="12" w:space="24" w:color="auto"/>
            </w:pgBorders>
            '''
            pgBorders = parse_xml(pgBorders_xml)
            sectPr.append(pgBorders)
    
    # 보고서 헤더 테이블
    header_table = doc.add_table(rows=6, cols=4)
    header_table.style = 'Table Grid'
    
    # 제목 행
    title_cell = header_table.rows[0].cells[0]
    title_cell.merge(header_table.rows[0].cells[3])
    title_cell.text = "⚡ 전력 사용량 분석 보고서"
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para = title_cell.paragraphs[0]
    title_para.runs[0].font.size = Cm(0.8)
    title_para.runs[0].bold = True
    
    # 헤더 정보 입력
    header_data = [
        ("보고처", "⚡ 에너지관리팀", "보고서명", f"전력 분석 보고서 ({period_label})"),
        ("장소", "🏭 본사", "취급분류", "○기밀 ●보통"),
        ("작성일자", datetime.now().strftime("%Y년 %m월 %d일"), "작성자", "🤖 AI 분석시스템"),
        ("분석범위", "실시간 전력 데이터, 한전 요금체계", "", ""),
        ("데이터출처", "전력량계 실시간 데이터, 한국전력공사 요금체계", "", "")
    ]
    
    for i, (col1, val1, col2, val2) in enumerate(header_data, 1):
        header_table.rows[i].cells[0].text = col1
        header_table.rows[i].cells[1].text = val1
        if col2:
            header_table.rows[i].cells[2].text = col2
            header_table.rows[i].cells[3].text = val2
        else:
            header_table.rows[i].cells[1].merge(header_table.rows[i].cells[3])
    
    doc.add_paragraph()
    
    # 보고내용 제목
    content_title = doc.add_heading('📊 전력 사용 분석 결과', level=1)
    content_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 요약 정보 추가
    if not current_data.empty:
        summary_paragraph = doc.add_paragraph()
        summary_paragraph.add_run("📋 분석 요약\n").bold = True
        
        total_kwh = current_data["전력사용량(kWh)"].sum()
        total_cost = current_data["전기요금(원)"].sum()
        avg_pf_day = current_data["지상역률(%)"].mean() if "지상역률(%)" in current_data.columns else 0
        avg_pf_night = current_data["진상역률(%)"].mean() if "진상역률(%)" in current_data.columns else 0
        
        summary_text = f"""
        • 분석 기간: {period_label}
        • 총 전력사용량: {total_kwh:,.1f} kWh
        • 총 전기요금: {total_cost:,.0f} 원
        • 평균 지상역률: {avg_pf_day:.1f}%
        • 평균 진상역률: {avg_pf_night:.1f}%
        """
        summary_paragraph.add_run(summary_text)
    
    # 차트 섹션 추가
    doc.add_heading('📈 차트 분석', level=2)
    
    # 월별 데이터 차트 생성
    try:
        if view_type == "월별" and not df.empty:
            monthly_data = df.groupby("년월").agg({
                "전력사용량(kWh)": "sum",
                "전기요금(원)": "sum"
            }).reset_index()
            
            chart_data = pd.DataFrame({
                "월": [str(x) for x in monthly_data["년월"]],
                "전력사용량": monthly_data["전력사용량(kWh)"],
                "전기요금": monthly_data["전기요금(원)"]
            })
            
            # matplotlib 차트 생성 및 문서에 추가
            chart_buffer = create_matplotlib_chart(
                chart_data[["월", "전력사용량"]], 
                "bar", 
                "월별 전력사용량 현황", 
                "월", 
                "전력사용량 (kWh)"
            )
            
            doc.add_paragraph("월별 전력사용량 추이:")
            doc.add_picture(chart_buffer, width=Inches(6))
            
    except Exception as e:
        doc.add_paragraph(f"차트 생성 중 오류 발생: {str(e)}")
    
    # 테이블 데이터 추가
    if not current_data.empty:
        doc.add_heading('📊 상세 데이터', level=2)
        
        # 요약 테이블
        summary_df = create_summary_table(current_data, period_label.replace("월", "").replace("기간", ""))
        
        table = doc.add_table(rows=1, cols=len(summary_df.columns))
        table.style = 'Table Grid'
        
        # 헤더 추가
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(summary_df.columns):
            hdr_cells[i].text = column
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # 데이터 추가
        for index, row in summary_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
    
    # 결론 및 권장사항
    doc.add_heading('💡 권장사항', level=2)
    recommendations = doc.add_paragraph()
    recommendations.add_run("에너지 효율 개선을 위한 권장사항:\n").bold = True
    
    rec_text = """
    1. 역률 관리: 지상역률 90% 이상, 진상역률 95% 이상 유지
    2. 피크 시간대 사용량 조절을 통한 요금 절감
    3. 정기적인 전력 사용 패턴 모니터링
    4. 에너지 효율 설비 도입 검토
    """
    recommendations.add_run(rec_text)
    
    # 보고서 완료
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run(f"보고서 생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}").italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    return doc

def main():
    # 메인 헤더
    st.markdown("""
    <div class="main-header">
        <h1>전력 사용량 분석 시스템</h1>
        <div class="subtitle">데이터 기반 에너지 효율 관리 및 분석 대시보드</div>
    </div>
    """, unsafe_allow_html=True)

    df = load_data()
    if df is None:
        st.stop()

    # 사이드바 설정
    with st.sidebar:
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-title">🔧 분석 설정</div>
        </div>
        """, unsafe_allow_html=True)
        
        # 분석 옵션
        st.markdown("#### 📊 상세 분석 옵션")
        numeric_columns = [
            "전력사용량(kWh)", "지상무효전력량(kVarh)", "진상무효전력량(kVarh)", 
            "탄소배출량(tCO2)", "지상역률(%)", "진상역률(%)", "전기요금(원)"
        ]
        
        # 데이터에 실제 존재하는 컬럼만 필터링
        available_columns = [col for col in numeric_columns if col in df.columns]
        
        if available_columns:
            col1_select = st.selectbox(
                "주요 분석 지표", 
                available_columns, 
                index=0,
                help="첫 번째 분석할 전력 지표를 선택하세요"
            )
            
            col2_select = st.selectbox(
                "비교 분석 지표", 
                available_columns, 
                index=min(1, len(available_columns)-1),
                help="두 번째 비교 분석할 전력 지표를 선택하세요"
            )
        else:
            st.error("분석 가능한 컬럼이 없습니다.")
            st.stop()

        st.markdown("---")
        st.markdown("#### 📄 보고서 생성")
        
        if st.button("📊 종합 분석 보고서 생성", type="primary"):
            with st.spinner("🔄 보고서 생성 중..."):
                try:
                    # 현재 설정 가져오기
                    view_type = st.session_state.get('analysis_period', '월별')
                    
                    if view_type == "월별":
                        selected_month = st.session_state.get('month_selector', 1)
                        current_year = df["년월"].dt.year.max()
                        current_data = df[
                            (df["년월"].dt.year == current_year) &
                            (df["년월"].dt.month == selected_month)
                        ]
                        period_label = f"{selected_month}월"
                    else:
                        selected_range = st.session_state.get('period_range_selector', None)
                        if selected_range and len(selected_range) == 2:
                            start_day, end_day = selected_range
                            current_data = df[
                                (df["날짜"] >= start_day) & 
                                (df["날짜"] <= end_day)
                            ]
                            period_label = f"{start_day} ~ {end_day} 기간"
                        else:
                            current_data = df
                            period_label = "전체 기간"
                    
                    # 최근 날짜 데이터
                    latest_date = df["날짜"].max()
                    daily_data = df[df["날짜"] == latest_date]
                    
                    # 보고서 생성
                    doc = create_comprehensive_docx_report_with_charts(
                        df, current_data, daily_data, latest_date, 
                        view_type, selected_month if view_type == "월별" else None, period_label
                    )
                    
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"⚡전력분석보고서_{timestamp}.docx"
                    
                    st.download_button(
                        label="📥 보고서 다운로드",
                        data=doc_buffer.getvalue(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
                    st.success("✅ 보고서 생성 완료!")
                    st.info(f"📁 파일명: {filename}")
                    
                except Exception as e:
                    st.error(f"❌ 보고서 생성 중 오류 발생: {str(e)}")

        st.markdown("<br>", unsafe_allow_html=True)
        
        # 시스템 정보
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-title">📊 시스템 정보</div>
            <div style="font-size: 0.85rem; color: #5f6368; line-height: 1.5;">
                • 분석 모델: 통계적 분석<br>
                • 데이터 범위: 전체 기간<br>
                • 분석 기법: 추세/패턴 분석<br>
                • 출력 형식: 대시보드/보고서
            </div>
        </div>
        """, unsafe_allow_html=True)

    # 메인 컨텐츠 영역 시작
    filtered_df = df.copy()
    summary_data = filtered_df.copy()
    period_label = "전체"

    # 필터링 옵션
    filter_col1, filter_col2, filter_col3, filter_col4 = st.columns([0.8, 0.5, 1, 0.7])
    
    with filter_col1:
        view_type = st.selectbox(
            "📅 분석 기간 설정", 
            ["월별", "일별"], 
            key="analysis_period",
            help="분석할 시간 단위를 선택하세요"
        )
    
    with filter_col2:
        if view_type == "월별":
            current_year = filtered_df["년월"].dt.year.max()
            months = list(range(1, 13))
            default_month = filtered_df["년월"].dt.month.max()
            selected_month = st.selectbox(
                "🗓️ 월", 
                months, 
                index=int(default_month) - 1, 
                key="month_selector",
                help="분석할 월을 선택하세요"
            )
        else:
            st.markdown("")
    
    with filter_col3:
        if view_type == "일별":
            from datetime import date
            date_range = (df["날짜"].min(), df["날짜"].max())
            selected_range = st.date_input(
                "📅 분석 기간", 
                value=(date(2024, 1, 1), date(2024, 1, 5)),
                min_value=date_range[0] if len(date_range) == 2 else df["날짜"].min(),
                max_value=date_range[1] if len(date_range) == 2 else df["날짜"].max(),
                key="period_range_selector",
                help="분석할 시작일과 종료일을 선택하세요"
            )
        else:
            st.markdown("")
    
    with filter_col4:
        st.markdown("")
    
    # 데이터 처리 로직
    current_data = pd.DataFrame()
    previous_data = pd.DataFrame()

    if view_type == "월별":
        current_data = filtered_df[
            (filtered_df["년월"].dt.year == current_year) & 
            (filtered_df["년월"].dt.month == selected_month)
        ]
        summary_data = current_data
        period_label = f"{selected_month}월"

        if selected_month > 1:
            prev_year = current_year
            prev_month = selected_month - 1
        else:
            prev_year = current_year - 1
            prev_month = 12
        previous_data = filtered_df[
            (filtered_df["년월"].dt.year == prev_year) & 
            (filtered_df["년월"].dt.month == prev_month)
        ]

    else:  # 일별
        if not isinstance(selected_range, tuple) or len(selected_range) != 2:
            st.warning("📅 날짜 범위를 선택해주세요")
        else:
            start_day, end_day = selected_range
            if start_day > end_day:
                st.warning("⛔ 시작 날짜가 종료 날짜보다 이후입니다.")
            else:
                period_df = filtered_df[
                    (filtered_df["날짜"] >= start_day) & 
                    (filtered_df["날짜"] <= end_day)
                ]
                if period_df.empty:
                    st.info(f"📊 {start_day} ~ {end_day} 구간에는 데이터가 없습니다.")
                else:
                    current_data = period_df
                    summary_data = period_df
                    period_label = f"{start_day} ~ {end_day} 기간"
                    
                    days = (end_day - start_day).days + 1
                    prev_start = start_day - timedelta(days=days)
                    prev_end = start_day - timedelta(days=1)
                    previous_data = filtered_df[
                        (filtered_df["날짜"] >= prev_start) & 
                        (filtered_df["날짜"] <= prev_end)
                    ]

    # 주요 지표 카드 표시
    if not summary_data.empty:
        main_metrics_card = create_main_metrics_card(summary_data, period_label)
        st.markdown(main_metrics_card, unsafe_allow_html=True)

    # 차트 섹션
    if view_type == "월별":
        st.markdown('<div class="section-header"><h2>📈 월별 전력 사용 트렌드 분석</h2></div>', unsafe_allow_html=True)
        
        # 월별 데이터 집계 시 컬럼 존재 확인
        if col1_select in filtered_df.columns and col2_select in filtered_df.columns:
            monthly_data = filtered_df.groupby("년월").agg({
                col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
                col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
            }).reset_index()
            monthly_data["년월_str"] = monthly_data["년월"].astype(str)

            fig = create_dual_axis_chart(
                monthly_data, "년월_str", col1_select, col2_select,
                f"월별 {col1_select} vs {col2_select} 비교 분석", 
                "월", col1_select, col2_select
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("선택한 분석 지표가 데이터에 존재하지 않습니다.")

    else:  # 일별
        if isinstance(current_data, pd.DataFrame) and not current_data.empty:
            st.markdown('<div class="section-header"><h2>📈 일별 전력 사용 트렌드 분석</h2></div>', unsafe_allow_html=True)
            
            if col1_select in current_data.columns and col2_select in current_data.columns:
                daily_data = current_data.groupby("날짜").agg({
                    col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
                    col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
                }).reset_index()

                fig = create_dual_axis_chart(
                    daily_data, "날짜", col1_select, col2_select,
                    f"{start_day} ~ {end_day} 날짜별 {col1_select} vs {col2_select}",
                    "날짜", col1_select, col2_select
                )
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.warning("선택한 분석 지표가 데이터에 존재하지 않습니다.")

    # 월별 분석일 때 비교 테이블
    if view_type == "월별" and not current_data.empty and not previous_data.empty:
        st.markdown('<div class="section-header"><h2>📊 전월 대비 상세 분석</h2></div>', unsafe_allow_html=True)
        comparison_df = create_comparison_table(current_data, previous_data, "월")
        if not comparison_df.empty:
            st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
            st.dataframe(comparison_df, use_container_width=True, hide_index=True)
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")

    # 특정일 시간별 분석
    st.markdown('<div class="section-header"><h2>🕐 시간별 에너지 사용 패턴 분석</h2></div>', unsafe_allow_html=True)

    col1, col2 = st.columns([3, 1])
    daily_df = pd.DataFrame()

    with col1:
        available_dates = sorted(filtered_df["날짜"].unique())
        if available_dates:
            min_d, max_d = available_dates[0], available_dates[-1]
            default_d = max_d
            selected_date = st.date_input(
                "📅 분석할 날짜 선택", 
                value=default_d, 
                min_value=min_d, 
                max_value=max_d, 
                key="daily_date_selector",
                help="시간별 상세 분석을 원하는 날짜를 선택하세요"
            )

            daily_df = filtered_df[filtered_df["날짜"] == selected_date]
            if daily_df.empty:
                st.warning(f"📊 {selected_date} 데이터가 없습니다.")
            else:
                # 시간별 데이터 집계 시 컬럼 존재 확인
                if col1_select in daily_df.columns and col2_select in daily_df.columns:
                    hourly_data = daily_df.groupby("시간").agg({
                        col1_select: ("sum" if col1_select not in ["지상역률(%)", "진상역률(%)"] else "mean"),
                        col2_select: ("sum" if col2_select not in ["지상역률(%)", "진상역률(%)"] else "mean")
                    }).reset_index()

                    # 24시간 전체 시간대 생성
                    full_hours = pd.DataFrame({"시간": list(range(24))})
                    hourly_data = pd.merge(full_hours, hourly_data, on="시간", how="left").fillna(0)

                    fig = create_dual_axis_chart(
                        hourly_data, "시간", col1_select, col2_select,
                        f"{selected_date} 시간별 {col1_select} vs {col2_select} 상세 분석",
                        "시간", col1_select, col2_select, add_time_zones=True
                    )

                    fig.update_xaxes(tickmode="linear", tick0=0, dtick=1, title_text="시간")
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.warning("선택한 분석 지표가 데이터에 존재하지 않습니다.")
        else:
            st.info("📊 선택된 조건에 맞는 데이터가 없습니다.")

    with col2:
        st.markdown("### 🚦 실시간 역률 모니터링")
        st.markdown("##### 한국전력공사 요금체계 기준")

        if available_dates and selected_date in available_dates:
            try:
                date_idx = available_dates.index(selected_date)
                if date_idx > 0:
                    previous_date = available_dates[date_idx - 1]
                    previous_daily_df = filtered_df[filtered_df["날짜"] == previous_date]

                    if not daily_df.empty and not previous_daily_df.empty:
                        # 역률 컬럼 존재 확인
                        if "지상역률(%)" in daily_df.columns and "진상역률(%)" in daily_df.columns:
                            # 주간 역률 계산 (09-23시)
                            current_daytime = daily_df[(daily_df['시간'] >= 9) & (daily_df['시간'] < 23)]
                            previous_daytime = previous_daily_df[(previous_daily_df['시간'] >= 9) & (previous_daily_df['시간'] < 23)]
                            # 야간 역률 계산 (23-09시)
                            current_nighttime = daily_df[(daily_df['시간'] >= 23) | (daily_df['시간'] < 9)]
                            previous_nighttime = previous_daily_df[(previous_daily_df['시간'] >= 23) | (previous_daily_df['시간'] < 9)]
                            
                            # 주간 지상역률 계산
                            if len(current_daytime) > 0:
                                current_daytime_raw = current_daytime['지상역률(%)'].mean()
                                current_daytime_pf = max(60, min(95, current_daytime_raw))
                            else:
                                current_daytime_pf = 90
                            
                            if len(previous_daytime) > 0:
                                previous_daytime_raw = previous_daytime['지상역률(%)'].mean()
                                previous_daytime_pf = max(60, min(95, previous_daytime_raw))
                            else:
                                previous_daytime_pf = 90
                            
                            # 야간 진상역률 계산
                            if len(current_nighttime) > 0:
                                current_leading_raw = current_nighttime['진상역률(%)'].mean()
                                if current_leading_raw > 0:
                                    current_nighttime_pf = max(60, current_leading_raw)
                                else:
                                    current_nighttime_pf = 100
                            else:
                                current_nighttime_pf = 100
                            
                            if len(previous_nighttime) > 0:
                                previous_leading_raw = previous_nighttime['진상역률(%)'].mean()
                                if previous_leading_raw > 0:
                                    previous_nighttime_pf = max(60, previous_leading_raw)
                                else:
                                    previous_nighttime_pf = 100
                            else:
                                previous_nighttime_pf = 100
                            
                            # 역률 카드 생성
                            daytime_card = create_simple_power_factor_card(
                                "주간", "🌅", current_daytime_pf, previous_daytime_pf, "daytime", "daytime-card"
                            )
                            nighttime_card = create_simple_power_factor_card(
                                "야간", "🌙", current_nighttime_pf, previous_nighttime_pf, "nighttime", "nighttime-card"
                            )
                            
                            st.markdown(daytime_card, unsafe_allow_html=True)
                            st.markdown(nighttime_card, unsafe_allow_html=True)
                        else:
                            st.info("📊 역률 데이터가 없습니다.")
                    else:
                        st.info("📊 선택된 날짜 또는 전일 데이터가 없습니다.")
                else:
                    if not daily_df.empty:
                        summary_df = create_summary_table(daily_df, "일")
                        st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
                        st.dataframe(summary_df, use_container_width=True, hide_index=True)
                        st.markdown("</div>", unsafe_allow_html=True)
                        st.info("📊 첫 번째 날짜로 전일 데이터가 없어 비교할 수 없습니다.")
                    else:
                        st.info("📊 선택된 날짜의 데이터가 없습니다.")
            except (ValueError, IndexError):
                st.info("📊 이전 날짜를 찾을 수 없습니다.")

    # 상세 비교 데이터 표
    if available_dates and selected_date in available_dates:
        try:
            date_idx = available_dates.index(selected_date)
            if date_idx > 0:
                previous_date = available_dates[date_idx - 1]
                previous_daily_df = filtered_df[filtered_df["날짜"] == previous_date]
                
                if not daily_df.empty and not previous_daily_df.empty:
                    st.markdown('<div class="section-header"><h2>📋 전일 대비 상세 비교 데이터</h2></div>', unsafe_allow_html=True)
                    comparison_df = create_comparison_table(daily_df, previous_daily_df, "일")
                    if not comparison_df.empty:
                        st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
                        st.dataframe(comparison_df, use_container_width=True, hide_index=True)
                        st.markdown("</div>", unsafe_allow_html=True)
        except (ValueError, IndexError):
            pass

    st.markdown("---")

    # 시간대별 현황 차트
    st.markdown('<div class="section-header"><h2>⚡ 작업유형별 전력 사용 현황</h2></div>', unsafe_allow_html=True)
    
    if not daily_df.empty:
        col_chart1, col_chart2 = st.columns([2, 1])
        with col_chart1:
            if "작업유형" in daily_df.columns and "전기요금(원)" in daily_df.columns:
                st.plotly_chart(create_hourly_stack_chart(daily_df), use_container_width=True)
            else:
                st.info("작업유형 또는 전기요금 데이터가 없습니다.")
        
        with col_chart2:
            if "작업유형" in daily_df.columns and "전력사용량(kWh)" in daily_df.columns:
                st.plotly_chart(create_concentric_donut_chart(daily_df), use_container_width=True)
            else:
                st.info("작업유형 또는 전력사용량 데이터가 없습니다.")

        # 작업유형별 상세 분석
        if "작업유형" in daily_df.columns:
            st.markdown('<div class="section-header"><h2>📊 작업유형별 상세 성과 분석</h2></div>', unsafe_allow_html=True)
            
            # 집계할 컬럼들 확인
            agg_dict = {}
            if "전력사용량(kWh)" in daily_df.columns:
                agg_dict["전력사용량_합계"] = ("전력사용량(kWh)", "sum")
            if "전기요금(원)" in daily_df.columns:
                agg_dict["전기요금_합계"] = ("전기요금(원)", "sum")
            if "지상역률(%)" in daily_df.columns:
                agg_dict["평균_지상역률"] = ("지상역률(%)", "mean")
            if "탄소배출량(tCO2)" in daily_df.columns:
                agg_dict["탄소배출량_합계"] = ("탄소배출량(tCO2)", "sum")
            
            if agg_dict:
                worktype_stats = daily_df.groupby("작업유형").agg(agg_dict).round(2)
                
                # 작업유형 한글명 매핑
                worktype_mapping = {
                    "Light_Load": "🌙 경부하",
                    "Medium_Load": "⚡ 중간부하", 
                    "Maximum_Load": "🔴 최대부하"
                }
                worktype_stats.index = [worktype_mapping.get(idx, idx) for idx in worktype_stats.index]
                
                st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
                st.dataframe(worktype_stats, use_container_width=True)
                st.markdown("</div>", unsafe_allow_html=True)
            else:
                st.info("작업유형별 분석에 필요한 데이터가 없습니다.")
        else:
            st.info("작업유형 데이터가 없습니다.")
    else:
        # 전체 데이터로 차트 생성
        col_chart1, col_chart2 = st.columns([2, 1])
        with col_chart1:
            if "작업유형" in filtered_df.columns and "전기요금(원)" in filtered_df.columns:
                st.plotly_chart(create_hourly_stack_chart(filtered_df), use_container_width=True)
            else:
                st.info("작업유형 또는 전기요금 데이터가 없습니다.")
        
        with col_chart2:
            if "작업유형" in filtered_df.columns and "전력사용량(kWh)" in filtered_df.columns:
                st.plotly_chart(create_concentric_donut_chart(filtered_df), use_container_width=True)
            else:
                st.info("작업유형 또는 전력사용량 데이터가 없습니다.")

        # 전체 작업유형별 상세 분석
        if "작업유형" in filtered_df.columns:
            st.markdown('<div class="section-header"><h2>📊 전체 작업유형별 상세 성과 분석</h2></div>', unsafe_allow_html=True)
            
            # 집계할 컬럼들 확인
            agg_dict = {}
            if "전력사용량(kWh)" in filtered_df.columns:
                agg_dict["전력사용량_합계"] = ("전력사용량(kWh)", "sum")
            if "전기요금(원)" in filtered_df.columns:
                agg_dict["전기요금_합계"] = ("전기요금(원)", "sum")
            if "지상역률(%)" in filtered_df.columns:
                agg_dict["평균_지상역률"] = ("지상역률(%)", "mean")
            if "탄소배출량(tCO2)" in filtered_df.columns:
                agg_dict["탄소배출량_합계"] = ("탄소배출량(tCO2)", "sum")
            
            if agg_dict:
                worktype_stats = filtered_df.groupby("작업유형").agg(agg_dict).round(2)
                
                # 작업유형 한글명 매핑
                worktype_mapping = {
                    "Light_Load": "🌙 경부하",
                    "Medium_Load": "⚡ 중간부하", 
                    "Maximum_Load": "🔴 최대부하"
                }
                worktype_stats.index = [worktype_mapping.get(idx, idx) for idx in worktype_stats.index]
                
                st.markdown('<div class="comparison-table">', unsafe_allow_html=True)
                st.dataframe(worktype_stats, use_container_width=True)
                st.markdown("</div>", unsafe_allow_html=True)
            else:
                st.info("작업유형별 분석에 필요한 데이터가 없습니다.")
        else:
            st.info("작업유형 데이터가 없습니다.")

    # 시스템 정보
    st.markdown("---")
    st.markdown('<div class="section-header"><h2>ℹ️ 시스템 정보</h2></div>', unsafe_allow_html=True)
    
    info_col1, info_col2, info_col3 = st.columns(3)
    
    with info_col1:
        st.markdown("""
        <div class="kpi-card">
            <div class="kpi-title">🏭 데이터 소스</div>
            <div style="font-size: 0.85rem; color: #5f6368; line-height: 1.5; text-align: left;">
                • LS 청주공장 실시간 전력량계<br>
                • 한국전력공사 요금체계<br>
                • 15분 단위 측정 데이터
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with info_col2:
        st.markdown("""
        <div class="kpi-card">
            <div class="kpi-title">📊 분석 범위</div>
            <div style="font-size: 0.85rem; color: #5f6368; line-height: 1.5; text-align: left;">
                • 전력사용량 및 요금 분석<br>
                • 역률 관리 및 최적화<br>
                • 탄소배출량 모니터링
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with info_col3:
        st.markdown("""
        <div class="kpi-card">
            <div class="kpi-title">🤖 시스템 기능</div>
            <div style="font-size: 0.85rem; color: #5f6368; line-height: 1.5; text-align: left;">
                • 실시간 패턴 분석<br>
                • 에너지 효율 최적화 제안<br>
                • 예측 기반 비용 절감
            </div>
        </div>
        """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()ㄴ